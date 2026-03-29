# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


OUT_PATH = Path(r"C:\Users\Ivan\Desktop\conspect_lecture_1_2_3_4_clean_ru.docx")


def add_term(doc: Document, term: str, definition: str, usage: str = "", code: str = "") -> None:
    doc.add_heading(term, level=2)
    doc.add_paragraph(f"Определение: {definition}")
    if usage:
        doc.add_paragraph(f"Применение в лабораторных: {usage}")
    if code:
        doc.add_paragraph("Пример:")
        code_paragraph = doc.add_paragraph()
        run = code_paragraph.add_run(code)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10)


def build_doc() -> Path:
    doc = Document()
    doc.add_heading("Конспект: Лекции 1, 2, 3, 4 (с примерами из лабораторных 1-3)", 0)
    doc.add_paragraph("Формат: терминология и краткая теория без формата вопросов и ответов.")

    doc.add_heading("Лекция 1. Клиент-серверная архитектура и основы HTML", 1)
    lecture_1 = [
        (
            "Сервер",
            "Компьютер или программа, предоставляющие ресурсы по запросу клиента.",
            "При публикации сайт отдает HTML/CSS/JS и изображения пользователю.",
            "",
        ),
        (
            "Клиент",
            "Программа или устройство, запрашивающее данные у сервера.",
            "В лабораторных роль клиента выполняет браузер.",
            "",
        ),
        (
            "Клиент-серверная архитектура",
            "Модель взаимодействия, в которой запрос и обработка разделены между клиентом и сервером.",
            "Все лабораторные проекты работают в этой модели.",
            "",
        ),
        (
            "HTML",
            "Язык разметки для описания структуры веб-страницы.",
            "Лаба 1 выполнена на HTML; в лабах 2-3 HTML является каркасом.",
            "<h1>Каталог ноутбуков</h1>",
        ),
        (
            "Тег",
            "Синтаксическая форма элемента в угловых скобках.",
            "Применяется во всех HTML-файлах проекта.",
            "<section>...</section>",
        ),
        (
            "Атрибут",
            "Дополнительный параметр элемента, задающий поведение или свойства.",
            "Используются href, src, alt, class, id, target, rel, data-*.",
            '<a href="items.html" target="_blank">Открыть</a>',
        ),
        (
            "Гиперссылка",
            "Переход к другому ресурсу по нажатию на элемент <a>.",
            "Меню и переходы на страницы товаров выполнены ссылками.",
            '<a href="products/asus.html">ASUS VivoBook</a>',
        ),
        (
            "Списки ul/ol/li",
            "ul — маркированный список, ol — нумерованный, li — пункт списка.",
            "Меню и характеристики товаров в лабах 1-3 реализованы списками.",
            "<ul>\n  <li>Процессор Intel Core i5</li>\n</ul>",
        ),
        (
            "Изображение img и alt",
            "img выводит изображение, alt задает альтернативный текст.",
            "Во всех лабораторных карточки товаров содержат изображения.",
            '<img src="assets/images/dell-thumb.png" alt="Ноутбук Dell">',
        ),
        (
            "Семантические теги",
            "Теги со смысловой ролью структуры: header, nav, main, section, article, footer, aside.",
            "Лабы 2-3 построены на семантической разметке.",
            "<header>...</header>\n<main>...</main>\n<footer>...</footer>",
        ),
    ]
    for row in lecture_1:
        add_term(doc, *row)

    doc.add_heading("Лекция 2. Основы CSS: синтаксис, каскад, свойства", 1)
    lecture_2 = [
        (
            "CSS",
            "Язык оформления, задающий внешний вид HTML-элементов.",
            "Лаба 2: styles/style.css; лаба 3: assets/css/style.css.",
            "body {\n  background-color: rgb(238, 243, 248);\n}",
        ),
        (
            "Селектор",
            "Часть правила CSS, выбирающая элементы для стилизации.",
            "Используются селекторы тегов, классов и id.",
            ".menu-list a { text-decoration: none; }",
        ),
        (
            "Блок объявлений",
            "Набор свойств и значений в фигурных скобках после селектора.",
            "Формирует итоговое оформление выбранного элемента.",
            ".product-card { border: 1px solid #ccc; padding: 18px; }",
        ),
        (
            "Каскад и наследование",
            "Каскад решает конфликты между правилами, наследование передает часть свойств потомкам.",
            "Более точное или более позднее правило может переопределить общее.",
            "body { color: rgb(33, 37, 41); }\n.action-link { color: rgb(100, 50, 100); }",
        ),
        (
            "Свойства текста и шрифта",
            "color, font-family, font-size, font-style, font-weight управляют типографикой.",
            "В лабе 2 эти свойства применены по требованиям задания.",
            ".short-description { font-size: 14px; font-style: italic; }",
        ),
        (
            "Отступы margin и padding",
            "margin задает внешний отступ, padding — внутренний.",
            "Ключевые свойства для расстояний между карточками, заголовками и блоками.",
            ".panel { padding: 26px; }\n.catalog-title { margin: 12px 0 8px; }",
        ),
        (
            "Сокращенная запись отступов",
            "1, 2, 3 или 4 значения задают стороны по правилам CSS.",
            "В стилях лабораторных используется регулярно.",
            "padding: 24px 0;\nmargin: 12px 0 8px;",
        ),
        (
            "RGB-цвета",
            "Способ задания цвета через интенсивность красного, зеленого и синего каналов.",
            "В лабораторной 2 RGB используется практически во всех правилах.",
            "color: rgb(60, 70, 90);",
        ),
    ]
    for row in lecture_2:
        add_term(doc, *row)

    doc.add_heading("Лекция 3. Продвинутые селекторы, комбинаторы и верстка", 1)
    lecture_3 = [
        (
            "Псевдокласс",
            "Селектор состояния элемента, например :hover или :focus.",
            "Лаба 2: hover для меню; лаба 3: focus-visible для интерактивных элементов.",
            ".menu-list a:hover { text-decoration: underline; }",
        ),
        (
            "Псевдоэлемент",
            "Селектор искусственной части элемента, например ::before и ::after.",
            "Лаба 3: декоративный маркер списка характеристик через ::before.",
            '.spec-list li::before { content: ""; }',
        ),
        (
            "Комбинаторы",
            "Способы связывать селекторы: потомок (пробел), дочерний (>), смежный (+), общий сосед (~).",
            "В проектах часто используется комбинатор потомка.",
            ".menu-list a { color: rgb(20, 60, 110); }",
        ),
        (
            "Специфичность селектора",
            "Вес селектора, влияющий на приоритет применения при конфликте правил.",
            "Общий селектор может быть переопределен более точным.",
            "body { color: #333; }\n.product-card p { color: #5a6478; }",
        ),
        (
            "Flexbox",
            "Модель раскладки для выравнивания элементов по одной оси.",
            "Лаба 2: блок с изображением и описанием; лаба 3: меню и группы кнопок.",
            ".product-hero { display: flex; flex-wrap: wrap; gap: 18px; }",
        ),
        (
            "Grid",
            "Двумерная модель раскладки для строк и колонок.",
            "Лаба 2: сетка каталога; лаба 3: основные сетки страниц.",
            ".catalog-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); }",
        ),
        (
            "@media",
            "Условное правило CSS для применения стилей при заданных параметрах экрана.",
            "В лабах 2-3 используется для адаптации под узкие экраны.",
            "@media (max-width: 640px) { .menu-list li { display: block; } }",
        ),
        (
            "Блочная и табличная верстка",
            "Блочная верстка строится на div/семантических блоках и CSS, табличная — на table.",
            "Лабораторные 1-3 выполнены блочной/семантической версткой.",
            "",
        ),
    ]
    for row in lecture_3:
        add_term(doc, *row)

    doc.add_heading("Лекция 4. Основы JavaScript и DOM", 1)
    lecture_4 = [
        (
            "JavaScript",
            "Язык программирования для добавления интерактивности веб-страницам.",
            "В лабораторной 3 JavaScript управляет фильтром каталога, корзиной и мобильным меню.",
            'const products = [{ id: 1, name: "ASUS VivoBook", price: 45990 }];',
        ),
        (
            "Где выполняется JavaScript",
            "Код выполняется в окружении: браузер, сервер (Node.js) или другое JS-окружение.",
            "В лабораторной 3 код выполняется в браузере.",
            '<script src="assets/js/script.js"></script>',
        ),
        (
            "Синтаксические конструкции",
            "Базовые части языка: инструкции, комментарии, операторы, переменные, функции, условия и циклы.",
            "В lab3 все эти конструкции используются в файле script.js.",
            "// комментарий\nconst total = cart.reduce((sum, item) => sum + item.price, 0);",
        ),
        (
            "Переменные var / let / const",
            "Средства объявления переменных: let — изменяемая, const — константная, var — устаревший способ.",
            "В lab3 применяются const и let.",
            "let cart = [];\nconst payButton = document.querySelector('#payButton');",
        ),
        (
            "Типы данных",
            "Основные типы: Number, String, Boolean, Null, Undefined, Array, Object.",
            "В lab3 используются числа (цены), строки (категории), массивы (cart), объекты (products).",
            "const product = { id: 1, name: 'ASUS', price: 45990, category: 'study' };",
        ),
        (
            "Оператор typeof",
            "Возвращает строку с типом переданного значения.",
            "Полезен при отладке логики фильтра и корзины.",
            "typeof 42; // 'number'",
        ),
        (
            "Операторы сравнения",
            "Сравнивают значения: >, <, >=, <=, ==, !=, ===, !==.",
            "В lab3 сравнение используется в условиях фильтрации и проверок.",
            "selectedCategory === 'all' || selectedCategory === card.dataset.category",
        ),
        (
            "Условное ветвление if / else if / else",
            "Выбирает выполняемый блок кода в зависимости от условия.",
            "В lab3 используется для проверки пустой корзины.",
            "if (cart.length === 0) {\n  alert('Корзина пуста');\n} else {\n  alert('Покупка прошла успешно');\n}",
        ),
        (
            "Тернарный оператор",
            "Краткая форма условия: условие ? значение1 : значение2.",
            "Используется в JS для компактных условных выражений.",
            "const text = cart.length ? 'Есть товары' : 'Пусто';",
        ),
        (
            "Циклы",
            "Конструкции for, while, for...of для многократного выполнения кода.",
            "В lab3 используется перебор массива и NodeList (forEach).",
            "cart.forEach((item) => console.log(item.name));",
        ),
        (
            "Функции",
            "Переиспользуемые блоки кода с параметрами и возможным return.",
            "В lab3 реализованы функции renderCart, addToCart, applyFilter и др.",
            "const calculateTotal = () => cart.reduce((total, item) => total + item.price, 0);",
        ),
        (
            "Function Declaration и Function Expression",
            "Два способа объявления функций: через function имя() и через присваивание функции в переменную.",
            "В lab3 чаще используются функциональные выражения через const.",
            "const createCartItem = (item, index) => { /* ... */ };",
        ),
        (
            "Стрелочные функции (arrow functions)",
            "Короткий синтаксис функций: (args) => expression или блок.",
            "В lab3 почти вся логика написана через стрелочные функции.",
            "addButtons.forEach((button) => {\n  button.addEventListener('click', () => { /* ... */ });\n});",
        ),
        (
            "Браузерное окружение",
            "Набор объектов и API, предоставляемых браузером поверх языка JavaScript.",
            "В lab3 используются document, window, alert, dataset, classList.",
            "window.addEventListener('resize', () => { /* ... */ });",
        ),
        (
            "DOM (Document Object Model)",
            "Объектная модель HTML-документа, где каждый элемент доступен как объект.",
            "В lab3 через DOM изменяются текст, классы и содержимое корзины.",
            "cartItems.innerHTML = \"\";\ncartItems.append(createCartItem(item, index));",
        ),
        (
            "Объект document",
            "Главная точка входа для работы со страницей в браузере.",
            "В lab3 document.querySelector* используется во всех ключевых частях кода.",
            "const categoryFilter = document.querySelector('#categoryFilter');",
        ),
        (
            "События и addEventListener",
            "Механизм реакции на действия пользователя (click, change и др.).",
            "В lab3 обработчики событий назначены кнопкам, фильтру и навигации.",
            "categoryFilter.addEventListener('change', applyFilter);",
        ),
        (
            "data-* и dataset",
            "Способ хранить пользовательские данные в HTML и получать их в JS.",
            "В lab3 у карточек есть data-id, data-price, data-category.",
            '<article class="product-card" data-id="1" data-price="45990" data-category="study">',
        ),
    ]
    for row in lecture_4:
        add_term(doc, *row)

    doc.add_heading("Приложение. Теория JavaScript из ПР-3, использованная в лабе 3", 1)
    js_terms = [
        (
            "DOM",
            "Объектная модель документа, позволяющая программно работать с HTML-элементами.",
            "В лабе 3 используется для корзины, фильтра и мобильного меню.",
            'const cartTotal = document.querySelector("#cartTotal");',
        ),
        (
            "querySelector / querySelectorAll",
            "Методы выборки элементов по CSS-селекторам.",
            "Применяются для кнопок, карточек и управляющих элементов формы.",
            'const addButtons = document.querySelectorAll(".add-to-cart");',
        ),
        (
            "События и addEventListener",
            "Механизм реакции на действия пользователя (click, change и т.д.).",
            "Используется для добавления в корзину, фильтра и кнопок оплаты/очистки.",
            'button.addEventListener("click", () => { /* ... */ });',
        ),
        (
            "data-* и dataset",
            "Способ хранить пользовательские данные в HTML и читать их в JS.",
            "В лабе 3 через data-id, data-price, data-category у карточек товара.",
            '<article class="product-card" data-id="1" data-price="45990" data-category="study">',
        ),
    ]
    for row in js_terms:
        add_term(doc, *row)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
